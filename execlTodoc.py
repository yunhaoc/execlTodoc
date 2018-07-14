#注意事项：1.execl中都是有效表格数据。2.行列都顶头写
#导入操作ececl和word文档的库
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
import xlrd

#打开execl
execl = xlrd.open_workbook('demoececl.xlsx')

#打开word文档
document = Document()

#计算execl中sheet的个数
sheetCount = len(execl.sheets())
print(sheetCount)
for index in range(0,sheetCount):
    table = execl.sheets()[index]
    tableRows = table.nrows
    tableCols = table.ncols
    wordTable = document.add_table(rows=tableRows, cols=tableCols,style="Table Grid")
    document.add_paragraph()
    for row in range(0,tableRows):
        for col in range(0,tableCols):
            data = table.cell(row,col)
            if data.value is not None:
                wordTable.cell(row,col).text = str(data.value)
document.save("test1.docx")
