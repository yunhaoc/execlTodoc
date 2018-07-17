"""
注意事项：1.execl中都是有效表格数据。2.行列都顶头写
导入操作ececl和word文档的库
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import xlrd
import sys


inputExeclPath = ""

"""
默认输出的word文档保存在程序运行的路径下
"""
outputWordPath = "testWord.docx"

"""
获取要转换的execl的名称(包含路径)和转换后保存的word文档的路径
"""


def get_input_args():
    """
    .检查是否有输入参数
    """
    global inputExeclPath
    global outputWordPath

    if len(sys.argv) == 1:
        print("please input the path of the execl:")
        return -1
    elif len(sys.argv) == 2:
        inputExeclPath = sys.argv[1]
    else:
        inputExeclPath = sys.argv[1]
        outputWordPath = sys.argv[2]


"""
brief:获取sheet的信息（sheet的行数和列数，合并单元格的信息）
args:
    sheetFd:传入sheet的标识符
    sheetInfo:传出sheet的行数和列数
    mergedInfo：传出合并单元格的信息
"""


def get_sheet_info(sheet_fd, sheet_info=None, merged_info=None):
    # 获取表格的行数和列数

    sheet_info.append(sheet_fd.nrows)
    sheet_info.append(sheet_fd.ncols)

    # 保存合并单元格的信息
    merged_dict = {'rs': 0, 're': 0, 'cs': 0, 'ce': 0}
    # 获取表格中合并的单元的信息(起始行，结束行，起始列，结束列)
    for crane in sheet_fd.merged_cells:
        temp = merged_dict.copy()
        temp['rs'], temp['re'], temp['cs'], temp['ce'] = crane
        """
         不能直接 merged_info.append(merged_dict)
        #因为因为字典d 是一个object ,而mergedInfo.append(merged_dict)并没有真正的将该字典在内存中再次创建。只是指向了相同的object。
         这也是python 提高性能，优化内存的考虑。
        """

        merged_info.append(temp)


def execl_to_word(e_fd, w_fd):
    """
    :param e_fd: excel文档的文件描述符
    :param w_fd: word文档的文件描述符
    :return:
    """
    # 计算execl中sheet的个数
    global write_re
    sheet_count = len(e_fd.sheets())
    # 若execl文档为空文档，则直接返回
    if sheet_count == 0:
        return

    table_index = 1
    for sheet in e_fd.sheets():
        row_and_col_info = []
        merged_info = []
        get_sheet_info(sheet, row_and_col_info, merged_info)

        i = 0
        write_re = 0
        while i < row_and_col_info[0]:
            is_exist = False
            if len(merged_info) > 0:
                for j in range(0, len(merged_info)):
                    if i == merged_info[j]['rs']:
                        merge_row = merged_info[j]['re'] - merged_info[j]['rs']
                        is_exist = True
                        break
            if not is_exist:
                merge_row = 1

            write_re = i + merge_row
            """
            添加表格上面的内容
            具体内容根据实际要求确定
            """
            paragraph_data = "    测试目的：" + str(sheet.cell(i, 0).value)
            w_fd.add_paragraph(paragraph_data)
            w_fd.add_paragraph("    测试平台：测试台架")
            w_fd.add_paragraph("    测试工具：XXXX")
            w_fd.add_paragraph()
            paragraph_data = "表" + str(table_index) + "-" + str(sheet.cell(i, 0).value)
            table_index = table_index + 1
            paragraph = w_fd.add_paragraph(paragraph_data)
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            word_table = w_fd.add_table(rows=merge_row, cols=row_and_col_info[1], style="Table Grid")
            w_fd.add_paragraph()

            # 每个表格开始写的时候行索引都为0
            row_index = 0
            # 向表格写数据
            for row in range(i, write_re):
                # 每行开始写数据时，列索引都为0
                col_index = 1

                for col in range(1, row_and_col_info[1]):
                    word_table.cell(row_index, 0).text = (str("步骤") + str(row_index + 1))
                    data = sheet.cell(row, col)
                    if data.value is not None:
                        try:
                            write_table_data = str("")
                            if isinstance(data.value, float):
                                if round(data.value) == data.value:
                                    write_table_data = str(round(data.value))
                                else:
                                    write_table_data = str(data.value)
                            else:
                                write_table_data = data.value
                            word_table.cell(row_index, col_index).text = str(write_table_data)
                            col_index = col_index + 1
                        except IndexError:
                            pass
                row_index = row_index + 1
            i = i + merge_row


def main():
    ret = get_input_args()
    if ret == -1:
        return

    # 打开execl
    try:
        execl_fd = xlrd.open_workbook(inputExeclPath)
    except FileNotFoundError:
        print("input the path of the execl!")
        return

    # 打开word文档
    document_fd = Document()
    execl_to_word(execl_fd, document_fd)
    document_fd.save(outputWordPath)


if __name__ == "__main__":
    main()
